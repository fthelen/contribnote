"""
Export module for CommentaryFlow.

Generates:
  - Word .docx with portfolio-specific letterhead
  - PDF (via docx2pdf if available)
  - commentary_sections.csv  (Snowflake-ready)
  - citations.csv             (Snowflake-ready)
"""

import csv
import json
import logging
from pathlib import Path
from datetime import datetime
from typing import Optional
from urllib.parse import urlparse

from . import db

logger = logging.getLogger(__name__)

OUTPUT_DIR = Path(__file__).parent / "output"
TEMPLATES_DIR = Path(__file__).parent / "templates"


# ---------------------------------------------------------------------------
# CSV Export (Snowflake-ready)
# ---------------------------------------------------------------------------

def export_snowflake_csvs(commentary_id: str) -> tuple[Path, Path]:
    """
    Write commentary_sections.csv and citations.csv to output/{commentary_id}/.
    Returns (sections_path, citations_path).
    """
    commentary = db.get_commentary(commentary_id)
    if not commentary:
        raise ValueError(f"Commentary {commentary_id} not found")

    folder = OUTPUT_DIR / commentary_id
    folder.mkdir(parents=True, exist_ok=True)

    sections = db.get_sections(commentary_id)
    all_citations = []
    for s in sections:
        cits = db.get_citations(commentary_id, s["section_key"], "gold")
        if not cits:
            cits = db.get_citations(commentary_id, s["section_key"], "silver")
        all_citations.extend(cits)

    # commentary_sections.csv
    sections_path = folder / "commentary_sections.csv"
    with open(sections_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=[
            "commentary_id", "portcode", "period_start", "period_end", "period_label",
            "section_key", "section_label", "section_type", "ticker", "security_name",
            "security_rank", "security_type", "contribution_to_return", "port_ending_weight",
            "gold_text", "bronze_text", "gold_word_count",
            "generated_at", "approved_at", "published_at", "generation_model",
        ])
        writer.writeheader()
        for s in sections:
            gold_text = s.get("gold_text") or s.get("silver_text") or s.get("bronze_text") or ""
            writer.writerow({
                "commentary_id": commentary_id,
                "portcode": commentary["portcode"],
                "period_start": commentary.get("period_start", ""),
                "period_end": commentary.get("period_end", ""),
                "period_label": commentary["period_label"],
                "section_key": s["section_key"],
                "section_label": s["section_label"],
                "section_type": s["section_type"],
                "ticker": s.get("ticker", ""),
                "security_name": s.get("security_name", ""),
                "security_rank": s.get("security_rank", ""),
                "security_type": s.get("security_type", ""),
                "contribution_to_return": s.get("contribution_to_return", ""),
                "port_ending_weight": s.get("port_ending_weight", ""),
                "gold_text": gold_text,
                "bronze_text": s.get("bronze_text", ""),
                "gold_word_count": len(gold_text.split()) if gold_text else 0,
                "generated_at": s.get("generated_at", ""),
                "approved_at": s.get("approved_at", ""),
                "published_at": commentary.get("published_at", ""),
                "generation_model": "",
            })

    # citations.csv
    citations_path = folder / "citations.csv"
    with open(citations_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=[
            "citation_id", "commentary_id", "portcode", "period_label",
            "section_key", "tier", "url", "title", "domain", "display_number",
            "source_origin", "removed_in_silver", "bronze_citation_id", "created_at",
        ])
        writer.writeheader()
        for cit in all_citations:
            writer.writerow({
                "citation_id": cit["citation_id"],
                "commentary_id": commentary_id,
                "portcode": commentary["portcode"],
                "period_label": commentary["period_label"],
                "section_key": cit["section_key"],
                "tier": cit["tier"],
                "url": cit["url"],
                "title": cit.get("title", ""),
                "domain": cit.get("domain", ""),
                "display_number": cit.get("display_number", ""),
                "source_origin": cit.get("source_origin", ""),
                "removed_in_silver": cit.get("removed_in_silver", 0),
                "bronze_citation_id": cit.get("bronze_citation_id", ""),
                "created_at": cit.get("created_at", ""),
            })

    logger.info(f"Exported CSVs for {commentary_id} → {folder}")
    return sections_path, citations_path


# ---------------------------------------------------------------------------
# Works Cited Assembly
# ---------------------------------------------------------------------------

def assemble_works_cited(commentary_id: str) -> list[dict]:
    """
    Deduplicate Gold citations across all sections, sorted by document order.
    Returns list of {number, title, domain, url}.
    """
    sections = db.get_sections(commentary_id)
    seen_urls: dict[str, int] = {}
    works_cited = []
    number = 1

    for section in sections:
        cits = db.get_citations(commentary_id, section["section_key"], "gold")
        if not cits:
            cits = db.get_citations(commentary_id, section["section_key"], "silver")
        for cit in cits:
            url = cit["url"]
            if url not in seen_urls:
                seen_urls[url] = number
                works_cited.append({
                    "number": number,
                    "title": cit.get("title") or url,
                    "domain": cit.get("domain") or _extract_domain(url),
                    "url": url,
                })
                number += 1

    return works_cited


# ---------------------------------------------------------------------------
# Word Document Export
# ---------------------------------------------------------------------------

def export_word(commentary_id: str) -> Path:
    """Generate Word document with portfolio letterhead."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    commentary = db.get_commentary(commentary_id)
    if not commentary:
        raise ValueError(f"Commentary {commentary_id} not found")

    portcode = commentary["portcode"]
    folder = OUTPUT_DIR / commentary_id
    folder.mkdir(parents=True, exist_ok=True)

    # Try portfolio-specific template, fall back to base
    template_path = TEMPLATES_DIR / "portfolios" / f"{portcode}.docx"
    base_path = TEMPLATES_DIR / "base_letterhead.docx"

    if template_path.exists():
        doc = Document(str(template_path))
    elif base_path.exists():
        doc = Document(str(base_path))
        logger.warning(f"No template for {portcode}, using base letterhead")
    else:
        doc = Document()
        _add_basic_letterhead(doc, commentary)

    # Populate content controls / placeholders
    _populate_document(doc, commentary, commentary_id)

    docx_path = folder / f"{commentary_id}.docx"
    doc.save(str(docx_path))
    logger.info(f"Word document saved: {docx_path}")
    return docx_path


def export_pdf(commentary_id: str) -> Path:
    """Generate PDF from Word document."""
    docx_path = export_word(commentary_id)
    pdf_path = docx_path.with_suffix(".pdf")

    try:
        import docx2pdf
        docx2pdf.convert(str(docx_path), str(pdf_path))
        logger.info(f"PDF saved: {pdf_path}")
        return pdf_path
    except ImportError:
        logger.warning("docx2pdf not installed — returning Word path instead")
        return docx_path
    except Exception as e:
        logger.error(f"PDF conversion failed: {e}")
        return docx_path


# ---------------------------------------------------------------------------
# Document population helpers
# ---------------------------------------------------------------------------

def _populate_document(doc, commentary: dict, commentary_id: str):
    """
    Replace placeholder strings in the document with actual content.
    Works with both content controls and simple {{PLACEHOLDER}} text.
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    sections = db.get_sections(commentary_id)
    works_cited = assemble_works_cited(commentary_id)

    replacements = _build_replacements(commentary, sections, works_cited)

    # Replace in all paragraphs and table cells
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)

    # If no overview placeholder was found, append content at end
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "{{OVERVIEW_TEXT}}" not in full_text and replacements.get("{{OVERVIEW_TEXT}}"):
        doc.add_heading("Overview", level=2)
        doc.add_paragraph(replacements["{{OVERVIEW_TEXT}}"])

        doc.add_heading("Securities", level=2)
        for s in sections:
            if s["section_type"] == "security":
                gold_text = s.get("gold_text") or s.get("silver_text") or s.get("bronze_text") or ""
                doc.add_heading(s["section_label"], level=3)
                doc.add_paragraph(gold_text)

        if replacements.get("{{OUTLOOK_TEXT}}"):
            doc.add_heading("Outlook", level=2)
            doc.add_paragraph(replacements["{{OUTLOOK_TEXT}}"])

        if works_cited:
            doc.add_heading("Works Cited", level=2)
            for wc in works_cited:
                doc.add_paragraph(
                    f'[{wc["number"]}]  {wc["title"]}. {wc["domain"]}. Retrieved from: {wc["url"]}'
                )


def _build_replacements(commentary: dict, sections: list[dict],
                         works_cited: list[dict]) -> dict:
    """Build the {{PLACEHOLDER}} → value mapping."""
    overview_text = ""
    outlook_text = ""
    security_lines = []

    for s in sections:
        gold_text = s.get("gold_text") or s.get("silver_text") or s.get("bronze_text") or ""
        if s["section_type"] == "overview":
            overview_text = gold_text
        elif s["section_type"] == "outlook":
            outlook_text = gold_text
        elif s["section_type"] == "security":
            security_lines.append(f"{s['section_label']}\n{gold_text}")

    works_cited_text = "\n".join(
        f'[{wc["number"]}]  {wc["title"]}. {wc["domain"]}. Retrieved from: {wc["url"]}'
        for wc in works_cited
    )

    approval_date = commentary.get("approved_at", "")[:10] if commentary.get("approved_at") else ""

    return {
        "{{PORTFOLIO_NAME}}": commentary["portcode"],
        "{{PERIOD_LABEL}}": commentary["period_label"],
        "{{OVERVIEW_TEXT}}": overview_text,
        "{{SECURITIES_TABLE}}": "\n\n".join(security_lines),
        "{{OUTLOOK_TEXT}}": outlook_text,
        "{{WORKS_CITED}}": works_cited_text,
        "{{APPROVAL_DATE}}": approval_date,
    }


def _replace_in_paragraph(para, replacements: dict):
    for placeholder, value in replacements.items():
        if placeholder in para.text:
            for run in para.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)


def _add_basic_letterhead(doc, commentary: dict):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    head = doc.add_heading(commentary["portcode"], level=1)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(commentary["period_label"])
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")


def _extract_domain(url: str) -> str:
    try:
        return urlparse(url).netloc.replace("www.", "")
    except Exception:
        return ""


# ---------------------------------------------------------------------------
# JSON save/load helpers (Bronze/Silver/Gold files)
# ---------------------------------------------------------------------------

def save_bronze_json(commentary_id: str, sections: list[dict]):
    _save_tier_json(commentary_id, "bronze", sections)


def save_silver_json(commentary_id: str, sections: list[dict]):
    _save_tier_json(commentary_id, "silver", sections)


def save_gold_json(commentary_id: str, sections: list[dict]):
    _save_tier_json(commentary_id, "gold", sections)


def _save_tier_json(commentary_id: str, tier: str, data):
    folder = OUTPUT_DIR / commentary_id
    folder.mkdir(parents=True, exist_ok=True)
    path = folder / f"{tier}.json"
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    logger.info(f"Saved {tier}.json for {commentary_id}")


def save_metadata_json(commentary_id: str, metadata: dict):
    folder = OUTPUT_DIR / commentary_id
    folder.mkdir(parents=True, exist_ok=True)
    path = folder / "metadata.json"
    with open(path, "w", encoding="utf-8") as f:
        json.dump(metadata, f, indent=2)
