"""
One-time script to create placeholder templates for CommentaryFlow.
Run:  python commentaryflow/create_templates.py

Creates:
  - commentaryflow/templates/base_letterhead.docx
  - commentaryflow/templates/portfolios/ABC.docx
  - commentaryflow/templates/portfolios/XYZ.docx
  - commentaryflow/surveys/survey_template.xlsx
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

TEMPLATES_DIR = Path(__file__).parent / "templates"
SURVEYS_DIR = Path(__file__).parent / "surveys"


def create_letterhead_docx(path: Path, portcode: str = None):
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("python-docx not installed. Run: pip install python-docx")
        return

    doc = Document()

    # Header
    header = doc.add_heading("", level=0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("FIRM NAME")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1a, 0x19, 0x17)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Portfolio Commentary").font.size = Pt(12)

    doc.add_paragraph()

    # Portfolio info placeholders
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Portfolio: {{PORTFOLIO_NAME}}  |  Period: {{PERIOD_LABEL}}").bold = True

    if portcode:
        pm_box = doc.add_paragraph()
        pm_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pm_box.add_run(f"[PM Photo placeholder — {portcode}]\nPortfolio Manager Name\nYears of Experience\nBio placeholder").font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    doc.add_paragraph()
    doc.add_heading("Overview", level=2)
    doc.add_paragraph("{{OVERVIEW_TEXT}}")

    doc.add_paragraph()
    doc.add_heading("Holdings", level=2)
    doc.add_paragraph("{{SECURITIES_TABLE}}")

    doc.add_paragraph()
    doc.add_heading("Outlook", level=2)
    doc.add_paragraph("{{OUTLOOK_TEXT}}")

    doc.add_paragraph()
    doc.add_heading("Works Cited", level=2)
    doc.add_paragraph("{{WORKS_CITED}}")

    # Footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Approved: {{APPROVAL_DATE}}  |  For professional use only. This commentary is confidential."
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    print(f"Created: {path}")


def create_survey_xlsx():
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        print("openpyxl not installed. Run: pip install openpyxl")
        return

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "PM Survey"

    ws.column_dimensions["A"].width = 50
    ws.column_dimensions["B"].width = 60

    header_fill = PatternFill("solid", fgColor="1a1917")
    header_font = Font(color="FFFFFF", bold=True)

    ws["A1"] = "Question"
    ws["B1"] = "PM Answer"
    for cell in [ws["A1"], ws["B1"]]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(wrap_text=True)

    questions = [
        "What were the primary drivers of portfolio performance this period?",
        "Which positions contributed most positively and why?",
        "Which positions detracted most and what is your assessment?",
        "What changes did you make to the portfolio this period and why?",
        "How has your sector/factor positioning changed?",
        "What is your outlook for the portfolio over the next quarter?",
        "Are there any macro themes or risks you are monitoring closely?",
        "Any other comments you would like included in the commentary?",
    ]

    for i, q in enumerate(questions, 2):
        ws[f"A{i}"] = q
        ws[f"A{i}"].alignment = Alignment(wrap_text=True)
        ws[f"B{i}"].alignment = Alignment(wrap_text=True)

    path = SURVEYS_DIR / "survey_template.xlsx"
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(path))
    print(f"Created: {path}")


if __name__ == "__main__":
    print("Creating CommentaryFlow templates…")
    create_letterhead_docx(TEMPLATES_DIR / "base_letterhead.docx")
    create_letterhead_docx(TEMPLATES_DIR / "portfolios" / "ABC.docx", portcode="ABC")
    create_letterhead_docx(TEMPLATES_DIR / "portfolios" / "XYZ.docx", portcode="XYZ")
    create_survey_xlsx()
    print("Done.")
